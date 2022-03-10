from docx import Document
from docx.shared import Inches

document = Document()
document.add_picture('me.png', width=Inches(1.0))

name = "Alex Salim"
phone = "+90 5522 6551 31"
email = "Marketer.Salim@gmail.com"

document.add_paragraph(name + ' | ' + email)
document.add_heading("About me:")
document.add_paragraph("I'm a marketer with +5 years of experience,\
 last Septemper I quit my job and started to trade in crypto market. \
The blockchain technology impressed me so much that I decided to change my carrer and \
start learning programming. \nMy ultimate goal is to build blockchains and work in (web3) field.")


document.add_heading("Education: ")
ba = "B.A in Marketing"
p = document.add_paragraph()
p.add_run(ba).bold = True
document.add_paragraph("   -Philadelphia University (2018)").italic = True
da = "Data Science Certificate"
p2 = document.add_paragraph()
p2.add_run(da).bold = True
document.add_paragraph("   -CodeCademey (2022)").italic = True

document.add_heading("Experience: ")
document.add_paragraph(
    "* Using Pyhton: I build a poject for U.S. Medical Insurance Costs\n    that does mutiple tasks such as:\n    -Calculating an estimated insurance cost.\n    -Show the difference between the real and estimated cost based on the person's input.\n    -Itreta through the data and make it readable for anyone.\n\n - You can check it out here: {https://github.com/Alex-salim}\n\n* Creat, edit and re-write file:\n    -CSV.\n    -Json.\n    -Docx: like this resume you're reading. I created it using Pyhton.")


document.save('resume.docx')
